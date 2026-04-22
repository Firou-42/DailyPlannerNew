import flet as ft
from core.task_manager import TaskManager
from datetime import date
from openpyxl import Workbook
from docx import Document


def build_export_screen(page: ft.Page, tm: TaskManager):
    page.clean()
    page.bgcolor = ft.Colors.GREY_100
    page.padding = 0

    content = ft.ListView(spacing=0, padding=0, expand=True)

    header = ft.Container(
        content=ft.Text("Экспорт отчёта", size=28, weight=ft.FontWeight.BOLD, color=ft.Colors.BLUE_GREY_800),
        padding=ft.padding.only(left=20, top=20, right=20, bottom=15),
    )
    content.controls.append(header)

    format_group = ft.RadioGroup(
        content=ft.Column([
            ft.Radio(value="word", label="Microsoft Word (.docx)"),
            ft.Radio(value="excel", label="Microsoft Excel (.xlsx)"),
        ]),
        value="word"
    )

    content.controls.append(ft.Container(
        content=ft.Text("Выберите формат:", size=16, color=ft.Colors.BLUE_GREY_800),
        padding=ft.padding.only(left=20, right=20, bottom=10),
    ))
    content.controls.append(ft.Container(content=format_group, padding=ft.padding.only(left=20, right=20)))

    content.controls.append(ft.Divider(height=1, color=ft.Colors.GREY_300))

    start_date_field = ft.TextField(label="Начало", value="2026-04-01", width=150)
    end_date_field = ft.TextField(label="Конец", value=date.today().isoformat(), width=150)

    content.controls.append(ft.Container(
        content=ft.Text("Период:", size=16, color=ft.Colors.BLUE_GREY_800),
        padding=ft.padding.only(left=20, top=15, bottom=10),
    ))
    content.controls.append(ft.Container(
        content=ft.Row([start_date_field, ft.Text("-", color=ft.Colors.GREY_600), end_date_field]),
        padding=ft.padding.only(left=20, right=20),
    ))

    content.controls.append(ft.Divider(height=1, color=ft.Colors.GREY_300))

    include_tasks = ft.Checkbox(label="Список задач", value=True)
    include_stats = ft.Checkbox(label="Статистика выполнения", value=True)
    include_categories = ft.Checkbox(label="Распределение по категориям", value=True)

    content.controls.append(ft.Container(
        content=ft.Text("Включить в отчёт:", size=16, color=ft.Colors.BLUE_GREY_800),
        padding=ft.padding.only(left=20, top=15, bottom=10),
    ))
    content.controls.append(ft.Container(content=include_tasks, padding=ft.padding.only(left=20, right=20)))
    content.controls.append(ft.Container(content=include_stats, padding=ft.padding.only(left=20, right=20)))
    content.controls.append(ft.Container(content=include_categories, padding=ft.padding.only(left=20, right=20)))

    def show_snack(message: str, is_error: bool = False):
        snack = ft.SnackBar(
            content=ft.Text(message, color=ft.Colors.WHITE),
            bgcolor=ft.Colors.RED if is_error else ft.Colors.GREEN,
        )
        page.overlay.append(snack)
        snack.open = True
        page.update()

    def export_to_excel(start_date: date, end_date: date):
        wb = Workbook()
        ws = wb.active
        ws.title = "Отчёт"

        if include_tasks.value:
            ws.append(["Дата", "Время", "Название", "Категория", "Приоритет", "Статус"])
            all_tasks = tm.get_all_tasks()
            priority_names = {1: "Низкий", 2: "Средний", 3: "Высокий"}

            for task in all_tasks:
                task_date_obj = date.fromisoformat(task[3])
                if start_date <= task_date_obj <= end_date:
                    status_text = "Выполнено" if task[8] == 1 else "Не выполнено"
                    ws.append([
                        task[3],
                        f"{task[4][:5]} - {task[5][:5]}",
                        task[1],
                        task[9],
                        priority_names.get(task[6], "Средний"),
                        status_text
                    ])

        if include_stats.value or include_categories.value:
            ws.append([])
            stats = tm.get_statistics()

            if include_stats.value:
                ws.append(["СТАТИСТИКА"])
                ws.append(["Всего задач", stats["total"]])
                ws.append(["Выполнено", stats["completed"]])
                ws.append(["Просрочено", stats["overdue"]])
                ws.append(["Процент выполнения", f"{stats['completion_rate']}%"])

            if include_categories.value:
                ws.append([])
                ws.append(["ПО КАТЕГОРИЯМ"])
                for cat_name, count in stats["by_category"].items():
                    ws.append([cat_name, count])

        filename = f"Отчёт_{start_date.isoformat()}_{end_date.isoformat()}.xlsx"
        wb.save(filename)
        return filename

    def export_to_word(start_date: date, end_date: date):
        doc = Document()
        doc.add_heading(f"Отчёт за период {start_date.isoformat()} - {end_date.isoformat()}", 0)

        if include_tasks.value:
            doc.add_heading("Список задач", level=1)
            all_tasks = tm.get_all_tasks()
            priority_names = {1: "Низкий", 2: "Средний", 3: "Высокий"}

            for task in all_tasks:
                task_date_obj = date.fromisoformat(task[3])
                if start_date <= task_date_obj <= end_date:
                    status_text = "Выполнено" if task[8] == 1 else "Не выполнено"
                    doc.add_paragraph(
                        f"{task[3]} {task[4][:5]}-{task[5][:5]} — {task[1]} ({task[9]}, {priority_names.get(task[6], 'Средний')}) — {status_text}"
                    )

        if include_stats.value:
            doc.add_heading("Статистика", level=1)
            stats = tm.get_statistics()
            doc.add_paragraph(f"Всего задач: {stats['total']}")
            doc.add_paragraph(f"Выполнено: {stats['completed']}")
            doc.add_paragraph(f"Просрочено: {stats['overdue']}")
            doc.add_paragraph(f"Процент выполнения: {stats['completion_rate']}%")

        if include_categories.value:
            doc.add_heading("По категориям", level=1)
            stats = tm.get_statistics()
            for cat_name, count in stats["by_category"].items():
                doc.add_paragraph(f"{cat_name}: {count}")

        filename = f"Отчёт_{start_date.isoformat()}_{end_date.isoformat()}.docx"
        doc.save(filename)
        return filename

    def export_clicked(e):
        try:
            start_date = date.fromisoformat(start_date_field.value)
            end_date = date.fromisoformat(end_date_field.value)
        except ValueError:
            show_snack("Неверный формат даты. Используйте ГГГГ-ММ-ДД", is_error=True)
            return

        if start_date > end_date:
            show_snack("Дата начала не может быть позже даты окончания", is_error=True)
            return

        try:
            if format_group.value == "excel":
                filename = export_to_excel(start_date, end_date)
            else:
                filename = export_to_word(start_date, end_date)
            show_snack(f"Файл сохранён: {filename}")
        except Exception as ex:
            show_snack(f"Ошибка при экспорте: {str(ex)}", is_error=True)

    content.controls.append(ft.Container(
        content=ft.ElevatedButton(
            "Экспортировать",
            icon=ft.Icons.DOWNLOAD,
            on_click=export_clicked,
            style=ft.ButtonStyle(bgcolor=ft.Colors.BLUE, color=ft.Colors.WHITE),
        ),
        padding=ft.padding.only(left=20, right=20, top=30, bottom=30),
    ))

    page.add(content)
    page.update()